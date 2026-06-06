"""Build report_final_new.docx with all correct numbers from current dataset."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMGS = os.path.join(ROOT, "scripts", "new_figures")
OUT  = os.path.join(ROOT, "report_final_new.docx")

IMG_LOGO        = os.path.join(ROOT, "scripts", "odt_images", "1000000000000258000002588A5CA7DF.png")
IMG_APP         = os.path.join(IMGS, "fig1_app.png")
IMG_BREAKDOWN   = os.path.join(IMGS, "fig2_breakdown.png")
IMG_CONVERGENCE = os.path.join(IMGS, "fig3_convergence.png")
IMG_SCHEDULE    = os.path.join(IMGS, "fig4_schedule.png")
IMG_TUNING      = os.path.join(IMGS, "fig5_tuning.png")

# ── Real numbers from current dataset ─────────────────────────────────────────
GEN0      = 15_050
GEN499    = 1_000
IMPROVE   = GEN0 - GEN499           # 14,050
PCT       = round(IMPROVE/GEN0*100, 1)  # 93.4%
BAD_TOT   = 201_650
HC1, HC2, HC3, SC1, SC2 = 105_000, 65_600, 27_500, 3_550, 0
TOP_A, TOP_B, TOP_SH = "COMP3320", "COMP3330", 36
POP50, POP200 = 950, 950
MUT002, MUT010 = 1_050, 950
MUT_PCT = round((MUT002 - MUT010) / MUT002 * 100, 1)   # 9.5%
SHOT_PEN = 1_150   # penalty shown in the 150-gen app screenshot

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def figure(img_path, width_in, cap_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(img_path, width=Inches(width_in))
    caption(cap_text)

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "999999")
    pBdr.append(bot)
    pPr.append(pBdr)

def cell(c, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if bg:
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)

GREY = "D9D9D9"
CTR  = WD_ALIGN_PARAGRAPH.CENTER

# ══ TITLE PAGE ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = CTR
p.paragraph_format.space_after = Pt(6)
p.add_run().add_picture(IMG_LOGO, width=Inches(1.4))

p = doc.add_paragraph()
p.alignment = CTR
r = p.add_run("Exam Timetable Scheduling Using Genetic Algorithm")
r.bold = True; r.font.size = Pt(22)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = CTR
p.paragraph_format.space_after = Pt(2)
p.add_run("COMP338 — Introduction to Artificial Intelligence").font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = CTR
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Adam Khabisa  |  1210475")
r.bold = True; r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = CTR
p.paragraph_format.space_after = Pt(2)
p.add_run("Instructor: Dr. Radi Jarrar").font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = CTR
p.paragraph_format.space_after = Pt(14)
r = p.add_run(
    "(couldn't find any partners so worked alone — "
    "I talked with Dr. Radi in person and he allowed me to work alone given my situation)"
)
r.font.size = Pt(9); r.italic = True

hline()

# ══ ABSTRACT ══════════════════════════════════════════════════════════════════
h1("Abstract")
body(
    f"I implemented a Genetic Algorithm (GA) to assign 22 course exams to 18 time "
    f"slots (6 days × 3 slots) for 95 students. The penalty function combines "
    f"three hard and two soft constraints. With default parameters (population 100, "
    f"mutation 0.05, tournament k=5), the penalty dropped from {GEN0:,} (initial "
    f"random best) to {GEN499:,} after 500 generations — a {PCT}% reduction. "
    f"Both small (pop=50) and large (pop=200) populations converge to {POP200:,} "
    f"in 300 generations; higher mutation (0.10) reaches {MUT010:,}."
)

# ══ 1. INTRODUCTION ═══════════════════════════════════════════════════════════
h1("1.  Introduction")
body(
    "Exam timetabling is NP‑hard. For a small instance of 22 courses and 95 "
    "students, the search space is 18²² possibilities. An exhaustive search is "
    "impossible. I built a GA from scratch (no GA libraries) as required by the "
    "project. The algorithm finds a good schedule by evolving a population of "
    "candidate solutions over many generations."
)

# ══ 2. PROBLEM DESCRIPTION ════════════════════════════════════════════════════
h1("2.  Problem Description")

h2("2.1  Data")
bullet("22 courses (from Course_Catalog).")
bullet("95 students, each enrolled in 4–7 courses (481 enrolments total).")
bullet("18 time slots: 6 days × 3 slots per day (09:00–11:00, 12:00–14:00, 15:00–17:00).")

h2("2.2  Chromosome")
body('A schedule is a Python dictionary: {course_code: slot_id}. For example:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
p.paragraph_format.space_after = Pt(5)
r = p.add_run('{"COMP2110": "D1S2", "COMP3320": "D3S1", ...}')
r.font.size = Pt(10); r.font.name = "Courier New"

h2("2.3  Penalty Function")
body("The fitness function returns the total penalty for a schedule — lower is better, zero is perfect.")

tbl = doc.add_table(rows=6, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["ID", "Type", "Rule", "Penalty"]):
    cell(tbl.rows[0].cells[i], h, bold=True, size=10, align=CTR, bg=GREY)
for ri, row in enumerate([
    ("HC1","Hard","Two courses in same slot with X shared students","shared_students × 1,000"),
    ("HC2","Hard","A student has >2 exams on one day","800 per student"),
    ("HC3","Hard","A student has ≥4 exams over two consecutive days","500 per student"),
    ("SC1","Soft","A student has exactly 2 exams on one day","50 per student"),
    ("SC2","Soft","Total exam days used > 5","100 per extra day"),
], start=1):
    for ci, v in enumerate(row):
        cell(tbl.rows[ri].cells[ci], v, size=10)
caption("Table 1.  Penalty rules applied by the fitness function.")

# ══ 3. GENETIC ALGORITHM DESIGN ═══════════════════════════════════════════════
h1("3.  Genetic Algorithm Design")

h2("3.1  Initialisation")
body("Each individual gets a random slot for every course. This creates a diverse starting population.")

h2("3.2  Selection — Tournament")
body("I pick 5 random individuals and select the one with the lowest penalty. I return a copy.")

h2("3.3  Crossover — Single-Point")
body(
    "Given two parents and a fixed order of courses, I split at a random point. "
    "The child takes the first part from parent A and the second part from parent B."
)

h2("3.4  Mutation")
body(
    "For each course, I change its slot with probability = mutation_rate. "
    "The new slot is chosen uniformly at random from all 18 slots."
)

h2("3.5  Elitism")
body(
    "I copy the two best individuals unchanged to the next generation. "
    "This guarantees the best penalty never gets worse across generations."
)

h2("3.6  Main GA Loop")
body("For each generation:")
for s in [
    "Compute penalty for all individuals.",
    "Update the global best if this generation improved.",
    "Keep the top two individuals unchanged (elitism).",
    "Fill the rest of the population by: selection → crossover → mutation.",
    "Replace the old population.",
]:
    bullet(s)

# ══ 4. IMPLEMENTATION ═════════════════════════════════════════════════════════
h1("4.  Implementation")
body(
    "I used Python 3.13, openpyxl (to read Excel), and matplotlib (plots). "
    "The GA logic is in core/evolve.py and does not depend on any GUI. "
    "The Tkinter UI (app.py) runs the GA in a background thread so the window "
    "stays responsive."
)
figure(IMG_APP, 5.4,
       f"Figure 1.  Exam Scheduler application window after a completed GA run "
       f"(150 generations, pop = 100, penalty = {SHOT_PEN:,}).")

# ══ 5. EXPERIMENTAL RESULTS ═══════════════════════════════════════════════════
h1("5.  Experimental Results")

h2("5.1  Fitness Test")
body(
    f"The provided Sample_Bad_Schedule gave a penalty of {BAD_TOT:,}. "
    f"This confirms the penalty function works as intended (e.g., {TOP_A} and "
    f"{TOP_B} share {TOP_SH} students and are in the same slot → "
    f"{TOP_SH * 1000:,} penalty)."
)
figure(IMG_BREAKDOWN, 5.4,
       f"Figure 2.  Penalty breakdown for the sample bad schedule (total = {BAD_TOT:,} points).")

h2("5.2  Default Run")
body("Parameters: pop=100, generations=500, mutation=0.05, k=5.")

tbl2 = doc.add_table(rows=4, cols=2)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Metric", "Value"]):
    cell(tbl2.rows[0].cells[i], h, bold=True, size=10, align=CTR, bg=GREY)
for ri, (m, v) in enumerate([
    ("Best penalty at generation 0",   f"{GEN0:,}"),
    ("Best penalty at generation 499", f"{GEN499:,}"),
    ("Total improvement",              f"{IMPROVE:,} pts ({PCT}%)"),
], start=1):
    cell(tbl2.rows[ri].cells[0], m, size=10)
    cell(tbl2.rows[ri].cells[1], v, size=10, align=CTR)
caption("Table 2.  Default-parameter GA run summary (500 generations).")

body(
    f"The penalty drops very steeply in the first 20–30 generations as the GA "
    f"eliminates HC1 clashes (same-slot conflicts). After generation ~50 the "
    f"curve flattens, resolving remaining HC2/HC3 day-load violations. "
    f"The final penalty of {GEN499:,} consists entirely of SC1 soft-constraint "
    f"violations (students with exactly 2 exams on the same day) — all hard "
    f"constraints are satisfied."
)
figure(IMG_CONVERGENCE, 5.6,
       f"Figure 3.  Best penalty per generation — default parameters, 500 generations.")

h2("Best Schedule")
body(
    f"The final schedule uses all 6 days. "
    f"The colour-coded grid below shows which exams share a slot."
)
figure(IMG_SCHEDULE, 6.0,
       f"Figure 4.  Colour-coded layout of the best schedule produced by the GA "
       f"(penalty = {GEN499:,}).")

# ══ 6. PARAMETER TUNING ═══════════════════════════════════════════════════════
h1("6.  Parameter Tuning")
body("I ran four extra tests for 300 generations each, changing one parameter at a time.")

tbl3 = doc.add_table(rows=5, cols=5)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Run","pop_size","mutation_rate","Generations","Final Penalty"]):
    cell(tbl3.rows[0].cells[i], h, bold=True, size=10, align=CTR, bg=GREY)
for ri, row in enumerate([
    ("1 — small population", "50",  "0.05", "300", f"{POP50:,}"),
    ("2 — large population", "200", "0.05", "300", f"{POP200:,}"),
    ("3 — low mutation",     "100", "0.02", "300", f"{MUT002:,}"),
    ("4 — high mutation",    "100", "0.10", "300", f"{MUT010:,}"),
], start=1):
    for ci, v in enumerate(row):
        cell(tbl3.rows[ri].cells[ci], v, size=10,
             align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else CTR)
caption("Table 3.  Parameter tuning results (300 generations, tournament k = 5 for all runs).")

figure(IMG_TUNING, 6.0,
       "Figure 5.  Convergence curves for four parameter configurations (300 generations each).")

h2("6.1  Effect of Population Size")
body(
    f"Both pop=50 and pop=200 converge to the same final penalty of {POP200:,} "
    f"in 300 generations. This shows the problem instance is not particularly "
    f"difficult: even a small population of 50 maintains enough diversity to "
    f"avoid premature convergence. The trade-off of larger population — longer "
    f"per-generation computation — yields no benefit here. Population size is "
    f"not the limiting factor for this 22-course, 95-student dataset."
)

h2("6.2  Effect of Mutation Rate")
body(
    f"A mutation rate of 0.10 (final penalty {MUT010:,}) outperforms 0.02 "
    f"(final penalty {MUT002:,}) — a {MUT_PCT}% improvement. Low mutation "
    f"allows the GA to exploit good partial solutions but can trap it in local "
    f"optima. Higher mutation injects fresh slot assignments that allow escape. "
    f"An excessively high rate (> 0.20) would degrade the algorithm toward "
    f"random search; the effective range for this problem appears to be "
    f"approximately 0.07–0.12."
)

# ══ 7. CONCLUSION ═════════════════════════════════════════════════════════════
h1("7.  Conclusion")
body(
    f"I built a Genetic Algorithm for university exam timetabling without using "
    f"any GA library. The algorithm lowered the penalty from {GEN0:,} (random "
    f"initial population) to {GEN499:,} after 500 generations with default "
    f"parameters — a {PCT}% drop. The final schedule satisfies all three hard "
    f"constraints (no same-slot clashes, no student with >2 exams per day, no "
    f"student with ≥4 exams over consecutive days); the remaining penalty of "
    f"{GEN499:,} comes entirely from SC1 soft violations."
)
body(
    f"Parameter tuning showed that population size has no effect on final quality "
    f"for this dataset (both 50 and 200 reach {POP200:,}), while a higher mutation "
    f"rate of 0.10 reduces the penalty from {MUT002:,} to {MUT010:,}. "
    f"The project runs with a Tkinter GUI that reads the Excel dataset on start-up "
    f"and saves best_schedule.txt, run_log.csv, and convergence.png to output/."
)
body(
    "Possible improvements: adaptive mutation that decreases over generations, "
    "more varied crossover methods (uniform or order-based), or a local search "
    "after the GA to swap slots and eliminate remaining SC1 violations."
)
body(
    "Note on libraries: The GA itself does not use any GA libraries — only "
    "openpyxl for reading Excel, matplotlib for plotting, and Tkinter for the interface."
)

# ══ REFERENCES ════════════════════════════════════════════════════════════════
h1("References")
for i, ref in enumerate([
    "Burke, E., & Petrovic, S. (2002). Recent Research Directions in Automated "
    "Timetabling. European Journal of Operational Research, 140(2), 266–280.",
    "Carter, M. W., Laporte, G., & Lee, S. Y. (1996). Examination Timetabling: "
    "Algorithmic Strategies and Applications. Journal of the Operational Research "
    "Society, 47(3), 373–383.",
    "Goldberg, D. E. (1989). Genetic Algorithms in Search, Optimization and "
    "Machine Learning. Addison-Wesley.",
    "Mitchell, M. (1996). An Introduction to Genetic Algorithms. MIT Press.",
    "Russell, S., & Norvig, P. (2021). Artificial Intelligence: A Modern Approach "
    "(4th ed.). Pearson.",
], start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after       = Pt(4)
    r = p.add_run(f"{i}.  {ref}")
    r.font.size = Pt(11)

doc.save(OUT)
print(f"Saved: {OUT}")
